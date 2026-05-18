from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Estilos
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

# Título principal
title = doc.add_paragraph()
title_run = title.add_run('Análisis Comparativo: Dos Estrategias de E-commerce Tabaquero')
title_run.font.size = Pt(26)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(31, 78, 120)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.paragraph_format.space_after = Pt(6)

# Subtítulo
subtitle = doc.add_paragraph('Smoke Jokers (WordPress + WooCommerce) vs. La Tabaquería (Jumpseller)')
subtitle.runs[0].font.italic = True
subtitle.runs[0].font.size = Pt(13)
subtitle.paragraph_format.space_after = Pt(18)

# RESUMEN EJECUTIVO
h1 = doc.add_heading('Resumen Ejecutivo', level=1)
h1.runs[0].font.color.rgb = RGBColor(31, 78, 120)

p = doc.add_paragraph()
p.add_run('Dos estrategias opuestas pero válidas para el mismo problema:')
p.paragraph_format.space_after = Pt(12)

# Tabla comparativa rápida
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'

hdr = table.rows[0].cells
hdr[0].text = 'Aspecto'
hdr[1].text = 'Smoke Jokers'
hdr[2].text = 'La Tabaquería'
for cell in hdr:
    cell.paragraphs[0].runs[0].bold = True

data = [
    ('Plataforma', 'WordPress + WooCommerce', 'Jumpseller'),
    ('Estilo', 'DIY (self-hosted)', 'SaaS/Cloud'),
    ('Stack', 'WordPress, PHP, MySQL', 'Jumpseller proprietary')
]

for i, row in enumerate(data, 1):
    table.rows[i].cells[0].text = row[0]
    table.rows[i].cells[1].text = row[1]
    table.rows[i].cells[2].text = row[2]

doc.add_paragraph()

# PLATAFORMA 1: SMOKE JOKERS
doc.add_heading('Plataforma 1: Smoke Jokers (WordPress + WooCommerce)', level=1).runs[0].font.color.rgb = RGBColor(31, 78, 120)

doc.add_heading('Stack Tecnológico', level=2)
for item in ['WordPress 6.9.4', 'WooCommerce 10.5.3', 'MySQL/MariaDB', 'Tema personalizado smoke_jokers', 'jQuery 3.7.1 + Bootstrap CSS']:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()

doc.add_heading('Ventajas', level=2)
advantages = [
    ('Control Total', 'Acceso a código, base de datos y servidor. Libertad absoluta para customizar.'),
    ('Costo Bajo Inicial', 'Hosting compartido $20-50/mes. WordPress y WooCommerce son gratuitos.'),
    ('Escalabilidad', 'Puede crecer desde 100 a 10.000+ productos sin cambiar de plataforma.'),
    ('SEO Nativo', 'WordPress es SEO-friendly. Documentado y probado.'),
    ('Ecosistema Maduro', 'Plugins para casi todo: membresías, puntos de retiro, age gate, analytics.'),
    ('Ingresos Recurrentes', 'WooCommerce Memberships implementado (Ser un Joker: $2.490/mes).'),
]
for title, desc in advantages:
    p = doc.add_paragraph()
    bold = p.add_run(title + ': ')
    bold.bold = True
    p.add_run(desc)
doc.add_paragraph()

doc.add_heading('Desventajas', level=2)
for item in [
    'JavaScript legacy (jQuery en 2026)',
    'Requiere mantenimiento técnico (actualizaciones, backups, seguridad)',
    'Sin soporte dedicado — depende de tu equipo técnico',
    'Costo crece con tráfico (servidor debe escalarse manualmente)',
    'Requiere conocimiento de WordPress para hacer cambios avanzados'
]:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()

# PLATAFORMA 2: LA TABAQUERÍA
doc.add_heading('Plataforma 2: La Tabaquería (Jumpseller)', level=1).runs[0].font.color.rgb = RGBColor(31, 78, 120)

doc.add_heading('Stack Tecnológico', level=2)
for item in ['Jumpseller (SaaS)', 'JavaScript + jQuery 3.4.1', 'Owl Carousel (sliders)', 'Almacenamiento en la nube', 'No hay acceso a código fuente']:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()

doc.add_heading('Ventajas', level=2)
advantages_j = [
    ('Zero Mantenimiento', 'Jumpseller gestiona servidor, backups, seguridad, actualizaciones.'),
    ('Tiempo Rápido al Mercado', 'Tienda online en horas, no semanas.'),
    ('Soporte Dedicado', 'Equipo de Jumpseller disponible.'),
    ('Escalabilidad Automática', 'Infraestructura crece sin intervención. Pagas solo por lo que usas.'),
    ('Costos Predecibles', 'Plan mensual fijo + comisión por venta. No hay sorpresas.'),
    ('Fácil de Usar', 'Dashboard intuitivo. No requiere técnico para cambios simples.'),
    ('Integraciones', 'Mercado Pago, PayPal, email, redes sociales integrados.'),
]
for title, desc in advantages_j:
    p = doc.add_paragraph()
    bold = p.add_run(title + ': ')
    bold.bold = True
    p.add_run(desc)
doc.add_paragraph()

doc.add_heading('Desventajas', level=2)
for item in [
    'Sin Control Total — No puedes acceder al código',
    'Limitado por la plataforma — Personalizaciones complejas requieren Jumpseller',
    'Costo Mensual — Incluso con 0 ventas',
    'Dependencia del Vendor — Si Jumpseller cierra o cambios de política, estás atrapado',
    'Lock-in — Migrar datos de Jumpseller es complicado',
    'Comisión por Transacción — 2-4% por venta (además del plan mensual)'
]:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()

# ANÁLISIS COMPARATIVO
doc.add_heading('Análisis Comparativo Detallado', level=1).runs[0].font.color.rgb = RGBColor(31, 78, 120)

# Tabla grande
comp_table = doc.add_table(rows=11, cols=3)
comp_table.style = 'Light Grid Accent 1'

hdr = comp_table.rows[0].cells
hdr[0].text = 'Métrica'
hdr[1].text = 'Smoke Jokers'
hdr[2].text = 'La Tabaquería'
for cell in hdr:
    cell.paragraphs[0].runs[0].bold = True

metrics = [
    ('Costo Inicial', '$500-1000 (dominio, hosting 1 año, tema)', '$99-299 (plan Jumpseller mensual)'),
    ('Costo Mensual', '$20-50', '$99-299+'),
    ('Comisión por Venta', '0%', '2-4%'),
    ('Control', 'Total', 'Limitado'),
    ('Mantenimiento', 'Tu equipo', 'Jumpseller'),
    ('Tiempo Setup', '2-4 semanas', '1-3 días'),
    ('Escalabilidad', 'Manual (servidor upgrade)', 'Automática'),
    ('SEO', 'Excelente (nativo)', 'Bueno (integrado)'),
    ('Customización', 'Ilimitada', 'Limitada a Jumpseller'),
    ('Riesgo', 'Bajo si hay soporte técnico', 'Medio (vendor lock-in)'),
]

for i, (metric, smoke, tabla) in enumerate(metrics, 1):
    comp_table.rows[i].cells[0].text = metric
    comp_table.rows[i].cells[1].text = smoke
    comp_table.rows[i].cells[2].text = tabla

doc.add_paragraph()

# PARA E-TABACO
doc.add_heading('Implicaciones para e-Tabaco (Monorepo 3 Marcas)', level=1).runs[0].font.color.rgb = RGBColor(31, 78, 120)

doc.add_heading('¿Qué Stack Elegir?', level=2)

p = doc.add_paragraph()
p.add_run('Tu objetivo es construir 3 ecommerce independientes con escalabilidad diferenciada.')
p.paragraph_format.space_after = Pt(6)

doc.add_heading('Opción A: Smoke Jokers (WordPress x3)', level=3)
doc.add_paragraph('Instalar WordPress + WooCommerce 3 veces', style='List Bullet')
doc.add_paragraph('Cada tienda completamente independiente', style='List Bullet')
doc.add_paragraph('Costo: $3 x ($20-50/mes) = $60-150/mes + hosting', style='List Bullet')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Resultado: ').bold = True
p.add_run('Control total, pero requiere 3x mantenimiento técnico.')
doc.add_paragraph()

doc.add_heading('Opción B: Jumpseller x3', level=3)
doc.add_paragraph('Crear 3 tiendas en Jumpseller', style='List Bullet')
doc.add_paragraph('Dashboard centralizado', style='List Bullet')
doc.add_paragraph('Costo: 3 x ($99-299/mes) = $297-897/mes + comisiones', style='List Bullet')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Resultado: ').bold = True
p.add_run('Fácil, sin mantenimiento, pero menos control y más caro.')
doc.add_paragraph()

doc.add_heading('Opción C: Django (Tu Stack)', level=3)
doc.add_paragraph('Monorepo centralizado con 3 sub-aplicaciones', style='List Bullet')
doc.add_paragraph('Una base de datos compartida (si quieres) o separadas (aisladas)', style='List Bullet')
doc.add_paragraph('Escalabilidad compartida pero flujos independientes', style='List Bullet')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Resultado: ').bold = True
p.add_run('Control total, escalabilidad moderna, mantenimiento centralizado. Tu ventaja competitiva.')
doc.add_paragraph()

# RECOMENDACIÓN FINAL
doc.add_heading('Recomendación Final', level=1).runs[0].font.color.rgb = RGBColor(31, 78, 120)

p = doc.add_paragraph()
bold = p.add_run('Para e-Tabaco: Django + PostgreSQL')
bold.bold = True
p.add_run(' (Opción C)')
p.paragraph_format.space_after = Pt(12)

doc.add_paragraph('Smoke Jokers y La Tabaquería prueban que el mercado funciona con cualquier stack. Pero hay razones para elegir Django:', style='List Bullet')
doc.add_paragraph()

for i, point in enumerate([
    ('Arquitectura Moderna', 'React/Vue frontend + API REST. Futuro-proof vs. jQuery legacy.'),
    ('Escalabilidad Centralizada', 'Monorepo = menos overhead que 3 instancias separadas.'),
    ('Control Total', 'Evitas vendor lock-in de Jumpseller o la complejidad de 3 WordPress.'),
    ('Datos Centralizados (opcional)', 'Puedes compartir catálogo, usuarios, analytics entre marcas si lo necesitas.'),
    ('Competencia en Frontend', 'La Tabaquería usa jQuery + Owl Carousel. Django moderno es ventaja competitiva.'),
], 1):
    p = doc.add_paragraph()
    bold = p.add_run(f'{i}. {point[0]}: ')
    bold.bold = True
    p.add_run(point[1])

doc.add_paragraph()

p = doc.add_paragraph()
italic = p.add_run('Bottom line: Smoke Jokers ganó con WordPress. La Tabaquería funciona con Jumpseller. Pero tu stack moderno te permite ganar con diferenciales que ellos no pueden.')
italic.italic = True

# Guardar
doc.save('/sessions/fervent-youthful-faraday/mnt/e-Tabaco/Analisis-Comparativo-Competencia.docx')
print("✓ Análisis comparativo creado")
