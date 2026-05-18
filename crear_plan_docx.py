from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Estilos
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

# Título
title = doc.add_paragraph()
title_run = title.add_run('Plan Estratégico: Back Office, Métricas y Escalabilidad')
title_run.font.size = Pt(26)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(31, 78, 120)

# Subtítulo
subtitle = doc.add_paragraph('e-Tabaco: De Frontend a Operaciones Completas')
subtitle.runs[0].font.italic = True
subtitle.runs[0].font.size = Pt(13)
subtitle.paragraph_format.space_after = Pt(18)

# RESUMEN
h1 = doc.add_heading('El Problema: Back Office Vacío', level=1)
h1.runs[0].font.color.rgb = RGBColor(31, 78, 120)

p = doc.add_paragraph()
p.add_run('Has enfocado e-Tabaco en el frontend (catálogo, carrito, checkout). Pero el ').font.color.rgb = RGBColor(0, 0, 0)
bold = p.add_run('back office administrativo está vacío')
bold.bold = True
p.add_run('. Sin esto, ¿cómo agregarás 200-500 productos por marca? ¿Cómo verás ventas? ¿Cómo lo delegas?')
p.paragraph_format.space_after = Pt(12)

# Estado Actual
doc.add_heading('Estado Actual vs. Lo Que Falta', level=1).runs[0].font.color.rgb = RGBColor(31, 78, 120)

doc.add_heading('✓ Lo que Existe', level=2)
for item in [
    'Django base configurado',
    'PostgreSQL lista',
    'Decisiones técnicas documentadas',
    'Propuesta de dominio + Deck de presentación'
]:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()

doc.add_heading('✗ Lo que NO Existe (Crítico)', level=2)
missing = [
    'Dashboard de métricas',
    'Gestión de productos (agregar/editar)',
    'Importador CSV (necesario para 500 productos)',
    'Gestión de órdenes',
    'Gestión de inventario',
    'Reportes de ventas',
    'Sistema de roles y permisos'
]
for item in missing:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item)
doc.add_paragraph()

# PRIORIDAD
doc.add_heading('Prioridad: Construir Back Office ANTES del Lanzamiento', level=1).runs[0].font.color.rgb = RGBColor(31, 78, 120)

p = doc.add_paragraph()
bold = p.add_run('¿Por qué?')
bold.bold = True
p.add_run(' Sin back office, no puedes:')
p.paragraph_format.space_after = Pt(6)

for item in [
    'Agregar 500 productos de forma eficiente',
    'Ver qué se vende (métricas)',
    'Gestionar stock',
    'Delegar a un asistente',
    'Operar el negocio después del lanzamiento'
]:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()

# Fases
doc.add_heading('Plan de Implementación (7 Semanas)', level=1).runs[0].font.color.rgb = RGBColor(31, 78, 120)

doc.add_heading('Fase 1: Backend MVP (Semanas 3-4) ⭐ CRÍTICA', level=2)
for item in [
    ('Autenticación', 'Django Admin + roles básicos'),
    ('Modelos de BD', 'Product, Order, Brand, Category'),
    ('Gestión de Productos', 'Formulario + Django Admin extendido'),
    ('Importador CSV', 'Para cargar 500 productos en 2 minutos'),
    ('Dashboard Mínimo', 'Ingresos hoy, últimas 10 órdenes, top productos')
]:
    p = doc.add_paragraph()
    bold = p.add_run(item[0] + ': ')
    bold.bold = True
    p.add_run(item[1])
doc.add_paragraph()

doc.add_heading('Fase 2: Frontend MVP (Semanas 1-2)', level=2)
for item in ['Landing page', 'Catálogo con filtros', 'Ficha de producto', 'Carrito básico', 'Gate de edad']:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()

doc.add_heading('Fase 3: Checkout + Pagos (Semanas 5-6)', level=2)
for item in ['Integración Mercado Pago', 'Página de checkout', 'Email de confirmación', 'Gestión de órdenes']:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()

# Respuesta a Dudas
doc.add_heading('Respuesta a Tus Dudas', level=1).runs[0].font.color.rgb = RGBColor(31, 78, 120)

doc.add_heading('❓ "¿Cómo veremos las métricas?"', level=2)
p = doc.add_paragraph()
bold = p.add_run('Fuentes de datos:')
bold.bold = True

for item in [
    ('Google Analytics (GTM)', 'Visitas, usuarios, tráfico. Ya está integrado.'),
    ('Dashboard Personalizado', 'Ingresos, órdenes, productos top. Lo construyes.'),
    ('Base de Datos', 'Reportes SQL: "¿Cuánto vendimos hoy?", "¿Top 5 productos?"'),
    ('Reportes en Excel', 'Descarga manual cada mes (ventas, márgenes, rentabilidad)')
]:
    p = doc.add_paragraph()
    bold = p.add_run(item[0] + ': ')
    bold.bold = True
    p.add_run(item[1])
doc.add_paragraph()

doc.add_heading('❓ "¿Cómo agregaremos 500 productos?"', level=2)
p = doc.add_paragraph()
bold = p.add_run('Respuesta: CSV + Importador')
bold.bold = True
p.paragraph_format.space_after = Pt(6)

steps = [
    'Editas catálogo en Excel (nombre, precio, marca, stock)',
    'Renombras imágenes: marca-tipo-peso.jpg',
    'En back office: click "Importar CSV" + sube archivo',
    'Sistema crea 500 productos en 2 minutos',
    'Revisa, corrige si hay errores, publica'
]

for i, step in enumerate(steps, 1):
    p = doc.add_paragraph(f'{i}. {step}')
doc.add_paragraph()

doc.add_heading('❓ "¿Cómo delegamos el back office?"', level=2)
p = doc.add_paragraph()
bold = p.add_run('Con Django Admin extendido + documentación simple')
bold.bold = True
p.paragraph_format.space_after = Pt(6)

for item in [
    'El asistente NO necesita saber código',
    'Aprende a usar Django Admin en 2-3 horas',
    'Solo puede hacer lo que le permites (permisos granulares)',
    'Documentación: manual de usuario (2 páginas) + plantilla Excel'
]:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()

doc.add_heading('❓ "¿Qué pasa con las 3 marcas?"', level=2)
p = doc.add_paragraph()
bold = p.add_run('Opción recomendada: Back office único con filtros por marca')
bold.bold = True
p.paragraph_format.space_after = Pt(6)

for item in [
    'Un dashboard centralizado',
    'Asistente 1 ve solo clubdeltabaco (por rol)',
    'Asistente 2 ve solo purotabaco (por rol)',
    'Asistente 3 ve solo zonatabaco (por rol)',
    'Reportes comparativos (las 3 marcas lado a lado)'
]:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()

# Arquitectura
doc.add_heading('Arquitectura Recomendada', level=1).runs[0].font.color.rgb = RGBColor(31, 78, 120)

p = doc.add_paragraph()
p.add_run('Todo en un monorepo Django, con estas capas:')
p.paragraph_format.space_after = Pt(12)

arch_items = [
    ('Frontend (Público)', 'Catálogo, carrito, checkout → es lo que ven los clientes'),
    ('Back Office (Admin)', 'Gestión de productos, órdenes, métricas → para ti y asistentes'),
    ('APIs (REST)', 'Endpoints para conectar con sistemas externos si es necesario'),
    ('Integraciones', 'Google Analytics, Mercado Pago, Email → todo integrado')
]

for title, desc in arch_items:
    p = doc.add_paragraph()
    bold = p.add_run(title + ': ')
    bold.bold = True
    p.add_run(desc)
doc.add_paragraph()

# Conclusión
doc.add_heading('Conclusión: Próximos Pasos', level=1).runs[0].font.color.rgb = RGBColor(31, 78, 120)

p = doc.add_paragraph()
bold = p.add_run('No estás atrasado.')
bold.bold = True
p.add_run(' Pero descuidar el back office desde el inicio es error común. Ahora que lo viste, arreglarlo cuesta 3-4 semanas, que es factible.')
p.paragraph_format.space_after = Pt(12)

doc.add_heading('Inmediato (esta semana)', level=2)
for item in [
    'Decidir: ¿Django Admin o Panel custom?',
    'Diseñar modelos de BD (Product, Order, Brand)',
    'Planificar importador CSV'
]:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()

doc.add_heading('Próximo (semanas 3-4)', level=2)
for item in [
    'Implementar back office MVP',
    'Cargar productos (CSV)',
    'Testear con ventas ficticias'
]:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()

doc.add_heading('Después (semanas 5+)', level=2)
for item in [
    'Checkout y pagos reales',
    'Dashboard de métricas',
    'Documentación para delegar'
]:
    doc.add_paragraph(item, style='List Bullet')

# Guardar
doc.save('/sessions/fervent-youthful-faraday/mnt/e-Tabaco/Plan-Back-Office-Visual.docx')
print("✓ Plan visual creado")
